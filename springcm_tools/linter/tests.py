from django.test import SimpleTestCase
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from xml.etree import ElementTree as ET
from .utils import lint

TYPE_UNORDERED = "1"
TYPE_ORDERED = "5"

def create_list(paragraph, type):
    p = paragraph._p #access to xml paragraph element
    pPr = p.get_or_add_pPr() #access paragraph properties
    numPr = OxmlElement('w:numPr') #create number properties element
    numId = OxmlElement('w:numId') #create numId element - sets bullet type
    numId.set(qn('w:val'), type) #set list type/indentation
    numPr.append(numId) #add bullet type to number properties list
    pPr.append(numPr) #add number properties to paragraph

def ms_wordify(input, ul_paragraphs=[], ol_paragraphs=[]):
    document = Document()
    for paragraph in input.split("\n"):
        document.add_paragraph(paragraph)

    for list in ul_paragraphs:
        p = document.paragraphs[list]
        p.style = "List Paragraph"
        create_list(p, TYPE_UNORDERED)

    for list in ol_paragraphs:
        p = document.paragraphs[list]
        p.style = "List Paragraph"
        create_list(p, TYPE_ORDERED)


    return document


class LintTests(SimpleTestCase):
    def test_unrecognized_tag(self):
        input = '<# <BadTagType Select="//Foo[text()=\'ok\']" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unrecognized tag type: 'BadTagType'")

    def test_one_tag_per_directive_pair(self):
        """There should be exactly one tag between directives"""
        input = '<# <Content Select="//Foo" Optional="true" /> <Content Select="//Foo" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Malformed XML")

    def test_nothing_between_directives(self):
        """There should not be anything other than a tag between directives"""
        input = '<# Hello <Content Select="//Foo" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Malformed XML")


class ParagraphLevelTests(SimpleTestCase):
    def test_unmatched_directives(self):
        """Test unmatched <# or #> directives"""
        input = '<# <Content Select="//Foo" Optional="true" />'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched #> or <# directive")

        input = '<Content Select="//Foo" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched #> or <# directive")

    def test_nested_directives(self):
        """Test nested <# <# #> #> directives"""
        input = '<# <# <Content Select="//Foo" Optional="true" /> #> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Nested #> or <# directives not allowed")

    def test_dont_parse_mergetags_if_paragraph_level_error(self):
        """Test that merge tags are not parsed if there's a paragraph-level error"""
        input = '<# <# <Content Select="//Foo"Optional="true" /> #> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Nested #> or <# directives not allowed")


class ContentTagTests(SimpleTestCase):
    def test_pass(self):
        """Confirm easy case passes"""
        input = '<# <Content Select="//Foo" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

    def test_self_closing_tags(self):
        """Test missing self-closing tag"""
        input = '<# <Content Select="//Foo" > #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Missing self-closing tag />")

    def test_malformed_xml(self):
        """Test malformed xml"""
        input = '<# <Content Select="//Foo"Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Malformed XML")

    def test_unrecognized_attributes(self):
        """Test unrecognized attributes"""
        input = '<# <Content Select="//Foo" Match="" /> #><# <Content Select="//Foo" Bar="" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 2)
        self.assertEqual(res[0].error, "Invalid attributes")
        self.assertEqual(res[1].error, "Invalid attributes")

    def test_required_attributes(self):
        """Test required attributes are present"""
        # I think only Select is required. TODO: investigate this
        input = '<# <Content Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <Content Select="//Foo" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

    def test_select(self):
        """Confirm the value of the Select attribute is valid XPath"""
        input = '<# <Content Select="\\badxpath" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Select attribute has invalid XPath")

    def test_optional(self):
        """Make sure the value of Optional is lowercase true or false"""
        input = '<# <Content Select="//Foo" Optional="True" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_track_name(self):
        """TrackName can contain only numbers, letters, spaces and periods, must start with a letter and cannot start with XML"""
        input = '<# <Content Select="//Foo" Optional="true" TrackName="Test Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

        input = '<# <Content Select="//Foo" Optional="true" TrackName="5est Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <Content Select="//Foo" Optional="true" TrackName="Te_t Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <Content Select="//Foo" Optional="true" TrackName="Te!t Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")


class TableRowTagTests(SimpleTestCase):
    def test_pass(self):
        """Confirm easy case passes"""
        input = '<# <TableRow Select="//Foo" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

    def test_self_closing_tags(self):
        """Test missing self-closing tag"""
        input = '<# <TableRow Select="//Foo" > #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Missing self-closing tag />")

    def test_malformed_xml(self):
        """Test malformed xml"""
        input = '<# <TableRow Select="//Foo"Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Malformed XML")

    def test_unrecognized_attributes(self):
        """Test unrecognized attributes"""
        input = '<# <TableRow Select="//Foo" Match="" /> #><# <TableRow Select="//Foo" Bar="" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 2)
        self.assertEqual(res[0].error, "Invalid attributes")
        self.assertEqual(res[1].error, "Invalid attributes")

    def test_required_attributes(self):
        """Test required attributes are present"""
        # I think only Select is required. TODO: investigate this
        input = '<# <TableRow Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <TableRow Select="//Foo" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

    def test_select(self):
        """Confirm the value of the Select attribute is valid XPath"""
        input = '<# <TableRow Select="\\badxpath" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Select attribute has invalid XPath")

    def test_optional(self):
        """Make sure the value of Optional is lowercase true or false"""
        input = '<# <TableRow Select="//Foo" Optional="True" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_track_name(self):
        """TrackName can contain only numbers, letters, spaces and periods, must start with a letter and cannot start with XML"""
        input = '<# <TableRow Select="//Foo" Optional="true" TrackName="Test Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

        input = '<# <TableRow Select="//Foo" Optional="true" TrackName="5est Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <TableRow Select="//Foo" Optional="true" TrackName="Te_t Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <TableRow Select="//Foo" Optional="true" TrackName="Te!t Track Name 123.456 XML" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")


class ConditionalTagTests(SimpleTestCase):
    def test_pass(self):
        """Confirm easy case passes"""
        input = '<# <Conditional Select="//Foo" Match="" /> #> Hello <# <EndConditional/> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

    def test_self_closing_tags(self):
        """Test missing self-closing tag"""
        input = '<# <Conditional Select="//Foo" Match=""> #> Hello <# <EndConditional/> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Missing self-closing tag />")

    def test_malformed_xml(self):
        """Test malformed xml"""
        input = '<# <Conditional Select="//Foo"Match="" /> #> Hello <# <EndConditional/> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Malformed XML")

    def test_unrecognized_attributes(self):
        """Test unrecognized attributes"""
        input = '<# <Conditional Select="//Foo" Bar="" Match="" /> #> Hello <# <EndConditional Match="" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 2)
        self.assertEqual(res[0].error, "Invalid attributes")
        self.assertEqual(res[1].error, "Invalid attributes")

    def test_required_attributes(self):
        """Test required attributes are present"""
        # I think either Select or Test are required. TODO: investigate this
        input = '<# <Conditional Match="" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_select(self):
        """Confirm the value of the Select attribute is valid XPath"""
        input = '<# <Conditional Select="\\badxpath" Match="" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Select attribute has invalid XPath")

    def test_test(self):
        """Make sure the value of Test is valid Xpath that returns true or false"""
        input = '<# <Conditional Test="true" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

        input = '<# <Conditional Test="\\badxpath" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Test attribute must be valid XPath that returns true or false")

    def test_select_and_test(self):
        """Make sure Select and Test attributes cant appear in the same tag"""
        input = '<# <Conditional Select="//Foo" Test="true" Match="" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_select_match(self):
        """Select attributes must have either Match or NotMatch but not both"""
        input = '<# <Conditional Select="//Foo" NotMatch="" Match="hello" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

        input = '<# <Conditional Select="//Foo" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_test_match(self):
        """Test attributes must not have either Match or NotMatch"""
        input = '<# <Conditional Test="true" Match="" /> #> Hello <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_unmatched_conditional_inline(self):
        """Inline Conditional must have matching inline EndConditional and vice-versa"""
        # No matching EndConditional
        input = '<# <Conditional Select="//Foo" Match="" /> #> Hello <# <Content Select="//Foo" Optional="true" /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched inline Conditional tag")

        # Extra EndConditional
        input = '<# <Conditional Select="//Foo" Match="" /> #> Hello <# <EndConditional /> #> Hello again <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched inline EndConditional tag")

        # Solo EndConditional (also should be on line by itself)
        input = 'Hello again <# <EndConditional /> #>'
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched inline EndConditional tag")

    def test_unmatched_conditional_paragraph(self):
        """Paragraph-level Conditional must have matching paragraph EndConditional and vice-versa"""
        # Passing
        para1 = '<# <Conditional Select="//Foo" Match="" /> #>'
        para2 = 'Hello <# <Content Select="//Foo" /> #> Hello again'
        para3 = '<# <EndConditional /> #>'
        input = '\n'.join([para1, para2, para3])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

        # Missing EndConditional
        para1 = '<# <Conditional Select="//Foo" Match="" /> #>'
        para2 = 'Hello <# <Content Select="//Foo" /> #> Hello again'
        input = '\n'.join([para1, para2])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched paragraph-level Conditional tag")

        # Missing Conditional
        para1 = 'Hello <# <Content Select="//Foo" /> #> Hello again'
        para2 = '<# <EndConditional /> #>'
        input = '\n'.join([para1, para2])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched paragraph-level EndConditional tag")

        # Bad EndConditional Position
        para1 = '<# <Conditional Select="//Foo" Match="" /> #>'
        para2 = 'Hello <# <Content Select="//Foo" /> #> Hello again <# <EndConditional /> #>'
        input = '\n'.join([para1, para2])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 2)
        self.assertEqual(res[0].error, "Unmatched paragraph-level Conditional tag")
        self.assertEqual(res[1].error, "Unmatched inline EndConditional tag")

        # Bad Conditional Position
        para1 = '<# <Conditional Select="//Foo" Match="" /> #> Hello <# <Content Select="//Foo" /> #>'
        para2 = '<# <EndConditional /> #>'
        input = '\n'.join([para1, para2])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 2)
        self.assertEqual(res[0].error, "Unmatched inline Conditional tag")
        self.assertEqual(res[1].error, "Unmatched paragraph-level EndConditional tag")

    def test_unmatched_conditional_complex(self):
        """Mixing and matching Conditional and EndConditional tags at various positions should work properly"""
        para1 = '<# <Conditional Select="//Foo" Match="" /> #>'
        para2 = 'Hello <# <Content Select="//Foo" /> #> <# <Conditional Select="//Foo" Match="sup" /> #> Hello again <# <EndConditional /> #>'
        para3 = '<# <EndConditional /> #>'
        input = '\n'.join([para1, para2, para3])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 0)

        para1 = '<# <Conditional Select="//Foo" Match="" /> #>'
        para2 = 'Hello <# <Content Select="//Foo" /> #> <# <Conditional Select="//Foo" Match="sup" /> #> Hello again <# <EndConditional /> #>'
        input = '\n'.join([para1, para2])
        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Unmatched paragraph-level Conditional tag")


class SuppressListItemTagTests(SimpleTestCase):
    def test_pass(self):
        """Confirm easy case passes"""
        input = '<# <SuppressListItem Select="//Foo" Match="" /> #> Hello'
        res = lint(ms_wordify(input, ul_paragraphs=[0]))
        self.assertEqual(len(res), 0)

    def test_unrecognized_attributes(self):
        """Test unrecognized attributes"""
        input = '<# <SuppressListItem Select="//Foo" Bar="" Match="" /> #> Hello'
        res = lint(ms_wordify(input, ul_paragraphs=[0]))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_required_attributes(self):
        """Test required attributes are present"""
        # I think either Select or Test are required. TODO: investigate this
        input = '<# <SuppressListItem Match="" /> #> Hello'
        res = lint(ms_wordify(input, ul_paragraphs=[0]))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "Invalid attributes")

    def test_not_in_list(self):
        """Test that it appears in a list item"""
        para1 = '<# <SuppressListItem Select="//Foo" Match="" /> #> Hello'
        para2 = 'This is no tags'
        para3 = '<# <SuppressListItem Select="//Foo" Match="" /> #> <# <Content Select="//Foo" Optional="true" /> #>'
        input = '\n'.join([para1, para2, para3])

        res = lint(ms_wordify(input))
        self.assertEqual(len(res), 2)
        self.assertEqual(res[0].error, "SuppressListItem must appear in a bullet or ordered list item")

        res = lint(ms_wordify(input, ul_paragraphs=[0,1,2]))
        self.assertEqual(len(res), 0)

        res = lint(ms_wordify(input, ol_paragraphs=[0,1,2]))
        self.assertEqual(len(res), 0)

        res = lint(ms_wordify(input, ol_paragraphs=[1,2]))
        self.assertEqual(len(res), 1)
        self.assertEqual(res[0].error, "SuppressListItem must appear in a bullet or ordered list item")



